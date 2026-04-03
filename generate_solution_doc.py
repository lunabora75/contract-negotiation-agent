from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# ── Colour palette ────────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1A, 0x2B, 0x4A)   # headings
ORANGE     = RGBColor(0xE8, 0x6C, 0x1E)   # accent / section rules
MID_GREY   = RGBColor(0x5A, 0x5A, 0x5A)   # body text
LIGHT_GREY = RGBColor(0xF5, 0xF5, 0xF5)   # table header fill
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

# ── Helper: shade a table cell ────────────────────────────────────────────────
def shade_cell(cell, hex_colour: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour)
    tcPr.append(shd)

# ── Helper: add a horizontal rule (paragraph border) ─────────────────────────
def add_hr(doc, colour_hex='E86C1E'):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), colour_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p

# ── Helper: heading 1 ─────────────────────────────────────────────────────────
def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = DARK_BLUE
    add_hr(doc)
    return p

# ── Helper: heading 2 ─────────────────────────────────────────────────────────
def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = ORANGE
    return p

# ── Helper: body paragraph ────────────────────────────────────────────────────
def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.font.size      = Pt(10)
            run.font.color.rgb = MID_GREY
            if i % 2 == 1:
                run.bold = True
    else:
        run = p.add_run(text)
        run.font.size      = Pt(10)
        run.font.color.rgb = MID_GREY
    return p

# ── Helper: bullet ────────────────────────────────────────────────────────────
def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    run = p.add_run(text)
    run.font.size      = Pt(10)
    run.font.color.rgb = MID_GREY
    return p

# ── Helper: simple table ──────────────────────────────────────────────────────
def add_table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # Header row
    hr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hr.cells[i]
        shade_cell(cell, '1A2B4A')
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = cell.paragraphs[0].add_run(h)
        run.bold            = True
        run.font.size       = Pt(9)
        run.font.color.rgb  = WHITE
        cell.paragraphs[0].paragraph_format.space_after  = Pt(3)
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    # Data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        fill = 'F5F5F5' if ri % 2 == 0 else 'FFFFFF'
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            shade_cell(cell, fill)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size      = Pt(9)
            run.font.color.rgb = MID_GREY
            cell.paragraphs[0].paragraph_format.space_after  = Pt(3)
            cell.paragraphs[0].paragraph_format.space_before = Pt(3)
    doc.add_paragraph()   # spacing after table
    return t

# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run('CONTRACT NEGOTIATION AGENT')
r.bold            = True
r.font.size       = Pt(24)
r.font.color.rgb  = DARK_BLUE

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run('High-Level Solution Approach Document')
r2.font.size      = Pt(14)
r2.font.color.rgb = ORANGE
r2.bold           = True

doc.add_paragraph()
add_hr(doc)

meta_lines = [
    ('Document Version', '1.0'),
    ('Date', datetime.date.today().strftime('%d %B %Y')),
    ('Status', 'Draft for Review'),
    ('Prepared by', 'Solution Architecture Team'),
    ('Classification', 'Internal / Confidential'),
]
for label, val in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_lbl = p.add_run(f'{label}:  ')
    r_lbl.bold           = True
    r_lbl.font.size      = Pt(10)
    r_lbl.font.color.rgb = DARK_BLUE
    r_val = p.add_run(val)
    r_val.font.size      = Pt(10)
    r_val.font.color.rgb = MID_GREY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '1. Executive Summary')
body(doc,
     'The Contract Negotiation Agent is an AI-powered platform that automates and '
     'optimises the negotiation of Statement of Work (SOW) contracts for professional '
     'IT services. By combining Claude (Anthropic\'s large language model) with '
     'proprietary market-rate benchmarks and a continuous-learning feedback loop, the '
     'system enables procurement and commercial teams to negotiate vendor contracts '
     'faster, more consistently, and at significantly lower cost than traditional '
     'manual processes.')

body(doc,
     'The platform operates as a digital commercial negotiator — analysing vendor '
     'submissions, positioning offers against market data, and engaging vendors in '
     'structured multi-turn negotiations — while keeping human approvers in control '
     'of the final decision.')

# ══════════════════════════════════════════════════════════════════════════════
#  2. BUSINESS PROBLEM
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '2. Business Problem Statement')

h2(doc, '2.1  The Challenge')
body(doc,
     'Enterprise organisations engage hundreds of IT vendors annually, each submitting '
     'bespoke SOW documents containing complex rate cards, payment schedules, liability '
     'clauses, and intellectual-property terms. Negotiating these contracts manually '
     'creates four interconnected problems:')

problems = [
    ('Cost Leakage',
     'Procurement teams lack real-time access to market-rate benchmarks, meaning '
     'vendor-proposed rates are often accepted with minimal challenge. Industry '
     'studies indicate that 60–70% of IT service contracts are signed at rates '
     '10–30% above market median — representing millions of pounds in avoidable spend.'),
    ('Inconsistency',
     'Negotiation outcomes depend heavily on the experience and availability of '
     'individual negotiators. Different buyers accept different rate levels for '
     'identical roles, creating inequitable spend profiles and audit risk.'),
    ('Capacity Constraints',
     'Commercial and legal teams are bottlenecks. Low-value or repeat-supplier '
     'contracts receive the same slow, resource-intensive treatment as strategic '
     'engagements, delaying project starts and frustrating stakeholders.'),
    ('Knowledge Loss',
     'Negotiation outcomes and lessons learned are captured informally (email, '
     'spreadsheets) or not at all. Each engagement effectively starts from scratch, '
     'with no institutional memory to guide future strategy.'),
]
for title, desc in problems:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f'{title}: ')
    r_t.bold           = True
    r_t.font.size      = Pt(10)
    r_t.font.color.rgb = DARK_BLUE
    r_d = p.add_run(desc)
    r_d.font.size      = Pt(10)
    r_d.font.color.rgb = MID_GREY

h2(doc, '2.2  Business Impact (Illustrative)')
add_table(doc,
    ['Metric', 'Current State', 'Target State', 'Potential Saving'],
    [
        ['Average rate vs. market median', '+18% above p50', 'At p50–p75', '10–18% rate reduction'],
        ['Time to first counter-offer', '3–5 business days', '< 30 minutes', '~97% reduction'],
        ['Contracts negotiated per FTE/month', '8–12', '40–60 (AI-assisted)', '4–5× throughput'],
        ['Negotiation outcome consistency', 'High variance', 'Standardised', 'Audit-ready'],
        ['Knowledge retention', 'Informal / lost', 'Structured & searchable', 'Continuous improvement'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  3. SOLUTION OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '3. Solution Overview')
body(doc,
     'The Contract Negotiation Agent is a full-stack web application that guides a '
     'contract through five automated stages, with a human approval gate at the end.')

stages = [
    ('Stage 1 — Document Ingestion',
     'The buyer uploads a vendor SOW (PDF, DOCX, or TXT). The platform extracts '
     'raw text and associates it with a unique session identifier.'),
    ('Stage 2 — AI Extraction',
     'Claude parses the SOW and populates a structured data model: client name, '
     'contract value, duration, individual roles with proposed hourly rates, payment '
     'terms, and key contractual clauses (liability, IP, termination).'),
    ('Stage 3 — Market Benchmarking',
     'Each extracted role is matched against an internal database of 50+ IT '
     'professional-service roles. The system calculates the proposed rate\'s '
     'percentile position (p25/p50/p75/p90), delta from market median, and '
     'monthly cost exposure at 160 hrs/month.'),
    ('Stage 4 — AI-Driven Negotiation',
     'Using the benchmark findings as evidence, the AI agent opens a negotiation '
     'dialogue directly with the vendor (or simulates one for buyer review). '
     'It applies a defined concession strategy: open at p50, concede to p75, '
     'concede to p90 for specialist roles only. All positions are cited with data.'),
    ('Stage 5 — Human Approval & Learning',
     'Once terms are agreed, the buyer\'s approver reviews the full transcript '
     'and approves or rejects the outcome. The decision and qualitative feedback '
     'are stored and injected into future sessions, enabling continuous improvement.'),
]
for title, desc in stages:
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_after = Pt(5)
    r_t = p.add_run(f'{title}:  ')
    r_t.bold           = True
    r_t.font.size      = Pt(10)
    r_t.font.color.rgb = ORANGE
    r_d = p.add_run(desc)
    r_d.font.size      = Pt(10)
    r_d.font.color.rgb = MID_GREY

# ══════════════════════════════════════════════════════════════════════════════
#  4. ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '4. Solution Architecture')

h2(doc, '4.1  Component Overview')
add_table(doc,
    ['Layer', 'Technology', 'Responsibility'],
    [
        ['Frontend', 'Next.js 15 / React 19 / TypeScript', 'Three-panel negotiation UI, file upload, session history, approval portal'],
        ['Backend API', 'FastAPI / Python / Uvicorn', 'REST API, session management, file parsing, orchestration'],
        ['AI Agent', 'Claude Sonnet (Anthropic)', 'SOW extraction, benchmark lookup, negotiation dialogue, feedback capture'],
        ['Data Store', 'In-memory + JSON file', 'Session state, market benchmarks, feedback/learning store'],
        ['Deployment', 'Railway (backend) / Vercel (frontend)', 'Cloud hosting, environment config, CORS management'],
    ]
)

h2(doc, '4.2  Data Flow')
flow_steps = [
    'User uploads SOW → FastAPI extracts text, creates session UUID',
    'NegotiationAgent.run_turn() called with SOW text',
    'Claude calls extract_sow_data tool → structured role/clause data returned',
    'Claude calls lookup_benchmarks tool → percentile analysis per role',
    'Claude generates counter-offer response based on benchmark evidence',
    'Multi-turn chat continues until negotiation is complete',
    'Buyer opens Approval Portal → reviews transcript → approves or rejects',
    'Feedback stored → injected into next session\'s system prompt',
]
for i, step in enumerate(flow_steps, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(f'  {i}.  {step}')
    r.font.size      = Pt(10)
    r.font.color.rgb = MID_GREY

doc.add_paragraph()

h2(doc, '4.3  AI Agent Design')
body(doc,
     'The core intelligence is a single NegotiationAgent built on Claude\'s tool-use '
     'capability. The agent is equipped with three tools:')
add_table(doc,
    ['Tool', 'Purpose', 'Outputs'],
    [
        ['extract_sow_data', 'Parse unstructured SOW text into structured contract data', 'Roles, rates, payment terms, clauses, confidence scores'],
        ['lookup_benchmarks', 'Compare proposed rates against internal market database', 'p25/p50/p75/p90, percentile, delta %, market position, monthly cost exposure'],
        ['save_feedback', 'Persist post-negotiation feedback for agent learning', 'Confirmation, cumulative feedback count'],
    ]
)

h2(doc, '4.4  Negotiation Strategy (Encoded in System Prompt)')
add_table(doc,
    ['Scenario', 'Opening Position', 'Maximum Concession'],
    [
        ['Standard IT roles', 'Market median (p50)', 'Experienced band (p75)'],
        ['Specialist / niche roles', 'Market median (p50)', 'Premium band (p90)'],
        ['Payment terms', 'Net 30', 'Net 45 (Net 60+ flagged as unacceptable)'],
        ['Liability cap', 'Contract value (1×)', 'Up to 2× contract value'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. KEY FEATURES
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '5. Key Features & Capabilities')

features = [
    ('Multi-format Document Ingestion', 'Supports PDF, DOCX, and TXT SOW files up to 10 MB. Text extraction handled server-side using pypdf and python-docx.'),
    ('Market Rate Benchmarking', 'Internal database of 50+ IT professional-service roles with p25/p50/p75/p90 hourly rate bands. Fuzzy-match logic handles variant role title naming.'),
    ('Real-time Negotiation Chat', 'Browser-based chat UI renders agent responses with full Markdown support — tables, bold emphasis, bullet lists — for clear, professional negotiation correspondence.'),
    ('Agent Learning Loop', 'Post-session feedback (rating 1–5, outcome, qualitative notes) is persisted and summarised into the agent\'s system prompt for subsequent sessions, enabling data-driven improvement.'),
    ('Buyer Approval Portal', 'Separate, shareable approval URL provides full negotiation transcript for review. Approver submits decision and comments; outcome is captured as feedback automatically.'),
    ('Session History', 'Browser localStorage maintains a history of recent negotiations, allowing buyers to revisit and compare outcomes.'),
    ('Risk Identification', 'The extraction stage identifies high-risk clauses (uncapped liability, missing termination-for-convenience, joint IP ownership, compliance obligations) and flags them to the buyer.'),
    ('Configurable AI Model', 'Model selection is environment-variable driven (MODEL env var), enabling model upgrades without code changes.'),
]
for title, desc in features:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    r_t = p.add_run(f'{title}: ')
    r_t.bold           = True
    r_t.font.size      = Pt(10)
    r_t.font.color.rgb = DARK_BLUE
    r_d = p.add_run(desc)
    r_d.font.size      = Pt(10)
    r_d.font.color.rgb = MID_GREY

# ══════════════════════════════════════════════════════════════════════════════
#  6. API REFERENCE
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '6. API Endpoints')
add_table(doc,
    ['Method', 'Endpoint', 'Purpose'],
    [
        ['GET',  '/api/health', 'Health check — returns status and active model'],
        ['POST', '/api/sessions', 'Create session: upload SOW, trigger initial AI extraction & opening offer'],
        ['POST', '/api/sessions/{id}/chat', 'Send a chat message; receive agent negotiation response'],
        ['GET',  '/api/sessions/{id}/summary', 'Retrieve full negotiation transcript and approval status'],
        ['POST', '/api/sessions/{id}/approval', 'Submit buyer approval or rejection decision'],
        ['POST', '/api/sessions/{id}/feedback', 'Submit qualitative feedback for agent learning'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  7. DEPLOYMENT
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '7. Deployment Architecture')
add_table(doc,
    ['Component', 'Platform', 'Config File', 'Key Environment Variables'],
    [
        ['Backend API', 'Railway.app (Docker)', 'backend/railway.json, Dockerfile', 'ANTHROPIC_API_KEY, MODEL, CORS_ORIGINS'],
        ['Frontend', 'Vercel (Next.js)', 'frontend/vercel.json', 'NEXT_PUBLIC_API_URL'],
        ['Data Persistence', 'JSON file (backend/data/store.json)', 'N/A', 'N/A'],
    ]
)

h2(doc, '7.1  Local Development')
for step in [
    'Copy .env.example → backend/.env and populate ANTHROPIC_API_KEY',
    'Install Python dependencies: pip install -r backend/requirements.txt',
    'Start backend: uvicorn main:app --reload  (port 8000)',
    'Install frontend dependencies: npm install  (frontend/)',
    'Start frontend: npm run dev  (port 3000)',
    'Navigate to http://localhost:3000 and upload a sample SOW',
]:
    bullet(doc, step)

# ══════════════════════════════════════════════════════════════════════════════
#  8. RISKS & CONSTRAINTS
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '8. Risks, Constraints & Mitigations')
add_table(doc,
    ['Risk / Constraint', 'Likelihood', 'Impact', 'Mitigation'],
    [
        ['In-memory session store lost on restart', 'High', 'Medium', 'Migrate to persistent DB (PostgreSQL / Redis) for production'],
        ['Benchmark data becomes stale', 'Medium', 'High', 'Establish quarterly benchmark refresh process; externalise to CMS'],
        ['Hallucinated extraction data from poor-quality PDFs', 'Medium', 'High', 'Add confidence-score thresholds; human review for low-confidence extractions'],
        ['Vendor adapts to agent negotiation patterns', 'Low', 'Medium', 'Randomise concession sequencing; update system prompt playbooks regularly'],
        ['Regulatory / legal review not automated', 'N/A', 'High', 'Agent flags risk clauses; human legal review remains mandatory'],
        ['Model API rate limits under high load', 'Low', 'Medium', 'Implement request queuing and retry logic; consider dedicated capacity'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  9. FUTURE ROADMAP
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '9. Future Roadmap')
add_table(doc,
    ['Phase', 'Initiative', 'Business Value'],
    [
        ['Phase 2', 'Persistent database (PostgreSQL)', 'Durable session history, audit trail, analytics'],
        ['Phase 2', 'Multi-agent orchestration (6-agent pattern)', 'Specialised agents for doc ingestion, risk, benchmarking, reporting'],
        ['Phase 2', 'Structured risk report export (PDF)', 'Formal risk assessment deliverable for legal/compliance'],
        ['Phase 3', 'ERP / P2P system integration', 'Auto-populate approved terms into procurement systems'],
        ['Phase 3', 'Live benchmark data feeds', 'Real-time rate data from market providers (e.g. Gartner, LinkedIn)'],
        ['Phase 3', 'Multi-supplier negotiation orchestration', 'Run parallel negotiations with competing vendors; select best outcome'],
        ['Phase 4', 'Analytics dashboard', 'Savings realised, win rate, average concession depth, role trends'],
        ['Phase 4', 'Vendor-facing portal', 'Vendors submit SOWs directly; agent negotiates in near-real-time'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  10. GLOSSARY
# ══════════════════════════════════════════════════════════════════════════════
h1(doc, '10. Glossary')
add_table(doc,
    ['Term', 'Definition'],
    [
        ['SOW', 'Statement of Work — a contract document defining the scope, deliverables, roles, rates, and terms of a professional services engagement'],
        ['p50 / Market Median', 'The hourly rate at which 50% of market participants charge less and 50% charge more'],
        ['p75', 'The rate at which 75% of market participants charge less — used as the standard maximum concession point'],
        ['p90', 'The rate at which 90% of market participants charge less — reserved for specialist/niche roles'],
        ['Tool Use', 'Anthropic\'s mechanism for Claude to call structured functions (tools) during a conversation'],
        ['Session', 'A single negotiation lifecycle: from SOW upload through to buyer approval'],
        ['Concession Ladder', 'The pre-defined sequence of rate offers an agent is authorised to make during negotiation'],
        ['Learning Loop', 'The feedback mechanism by which past negotiation outcomes improve future agent behaviour'],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  Save
# ══════════════════════════════════════════════════════════════════════════════
out_path = '/home/user/contract-negotiation-agent/Contract_Negotiation_Agent_Solution_Approach.docx'
doc.save(out_path)
print(f'Document saved to: {out_path}')
