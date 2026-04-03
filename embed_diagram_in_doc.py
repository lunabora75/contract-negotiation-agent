from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document('/home/user/contract-negotiation-agent/Contract_Negotiation_Agent_Solution_Approach.docx')

DARK_BLUE = RGBColor(0x1A, 0x2B, 0x4A)
ORANGE    = RGBColor(0xE8, 0x6C, 0x1E)
MID_GREY  = RGBColor(0x5A, 0x5A, 0x5A)

def add_hr(doc, colour_hex='E86C1E'):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), colour_hex)
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.space_before = Pt(0)
    return p

# Add page break then Architecture section
doc.add_page_break()

# Section heading
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(18)
p.paragraph_format.space_after  = Pt(4)
run = p.add_run('APPENDIX A — SOLUTION ARCHITECTURE DIAGRAM')
run.bold           = True
run.font.size      = Pt(14)
run.font.color.rgb = DARK_BLUE
add_hr(doc)

# Caption
cap = doc.add_paragraph()
cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
cap.paragraph_format.space_before = Pt(6)
r = cap.add_run('Figure 1: Contract Negotiation Agent — High-Level Solution Architecture')
r.font.size      = Pt(9)
r.font.color.rgb = MID_GREY
r.bold           = True
r.italic         = True

# Insert diagram
pic_para = doc.add_paragraph()
pic_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
pic_para.paragraph_format.space_before = Pt(6)
pic_para.paragraph_format.space_after  = Pt(8)
run = pic_para.add_run()
run.add_picture('/home/user/contract-negotiation-agent/Contract_Negotiation_Agent_Architecture.png',
                width=Inches(6.8))

# Key notes below diagram
p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(8)
r2 = p2.add_run('Diagram Notes')
r2.bold           = True
r2.font.size      = Pt(10)
r2.font.color.rgb = ORANGE

notes = [
    'All frontend-to-backend communication is via REST over HTTPS.',
    'The AI Agent communicates with the Anthropic API (Claude Sonnet) using the tool-use pattern — Claude calls structured tools (extract_sow_data, lookup_benchmarks, save_feedback) rather than generating free-form responses for data operations.',
    'Session state is currently held in-memory on the backend process; a persistent store (e.g. PostgreSQL/Redis) is recommended for production.',
    'The Feedback Store (store.json) is read at session start to inject the last 8 sessions\' learning context into the agent system prompt.',
    'Deployment targets: Backend → Railway.app (Docker), Frontend → Vercel (Next.js).',
]
for note in notes:
    p_note = doc.add_paragraph(style='List Bullet')
    p_note.paragraph_format.space_after = Pt(3)
    r_note = p_note.add_run(note)
    r_note.font.size      = Pt(9)
    r_note.font.color.rgb = MID_GREY

out = '/home/user/contract-negotiation-agent/Contract_Negotiation_Agent_Solution_Approach.docx'
doc.save(out)
print(f'Updated document saved: {out}')
